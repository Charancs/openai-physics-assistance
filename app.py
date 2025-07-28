import os
import json
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file
from werkzeug.utils import secure_filename
import openai
from dotenv import load_dotenv
from PIL import Image
import pytesseract
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Load environment variables
load_dotenv()

# Set up Flask app with explicit template and static folder paths
template_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'app', 'templates')
static_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'app', 'static')
app = Flask(__name__, template_folder=template_dir, static_folder=static_dir)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'app', 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure OpenAI
openai.api_key = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")

# Add proxy configuration (can be commented out for personal laptop use)
# Get proxies from environment or .env file
http_proxy = os.getenv("HTTP_PROXY", "http://proxy:8080")
https_proxy = os.getenv("HTTPS_PROXY", "http://proxy:8080")

# Flag to enable/disable proxy (set to False to disable)
USE_PROXY = True  # Set to False when using on a personal laptop without proxy

# Add timeout configuration for API requests
from openai import OpenAI
import httpx

# Configure proxies for httpx
proxies = None
if USE_PROXY:
    # Format proxies according to httpx expected format
    proxies = {
        "http://": http_proxy,
        "https://": https_proxy
    }
    print(f"Using proxies: {proxies}")
    # Set environment variables for OpenAI client
    os.environ["HTTP_PROXY"] = http_proxy
    os.environ["HTTPS_PROXY"] = https_proxy

# Create OpenAI client with proxy and timeout settings
openai_client = OpenAI(
    api_key=openai.api_key,
    timeout=httpx.Timeout(60.0, connect=10.0)
)

# Configure Tesseract OCR path (Windows default installation path)
pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_PATH", r'C:\Program Files\Tesseract-OCR\tesseract.exe')

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_image(image_path):
    """Extract text from the uploaded image using OCR"""
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from image: {e}")
        return None

def process_physics_question(question):
    """Process a physics question using OpenAI API"""
    try:
        # Add more detailed error logging
        print(f"Attempting to process question with OpenAI: {question[:50]}...")
        
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are a physics problem solver for JEE/NEET. Be concise. For each problem: "
                                             "1. Provide step-by-step numerical solution with proper units "
                                             "2. End with a shortcut technique/trick for similar problems. No explanatory fluff."},
                {"role": "user", "content": f"Solve: {question}"}
            ],
            temperature=0.1,
            max_tokens=2000
        )
        
        # Extract token usage information
        token_usage = {
            'prompt_tokens': response.usage.prompt_tokens,
            'completion_tokens': response.usage.completion_tokens,
            'total_tokens': response.usage.total_tokens
        }
        
        # Extract and return the assistant's response with token usage
        return {
            "content": response.choices[0].message.content,
            "token_usage": token_usage
        }
    
    except Exception as e:
        import traceback
        print(f"Error processing with OpenAI: {str(e)}")
        print(f"API Key configured: {'Yes' if openai.api_key else 'No'}")
        print(f"Model being used: {OPENAI_MODEL}")
        print(traceback.format_exc())
        
        # Return a more user-friendly error message
        error_msg = str(e)
        if "Connection" in error_msg:
            user_message = "Connection error: Could not connect to OpenAI. Please check your internet connection and try again."
        elif "API key" in error_msg:
            user_message = "API key error: Please check your OpenAI API key configuration."
        elif "not found" in error_msg:
            user_message = f"Model error: The model {OPENAI_MODEL} was not found. Please try another model."
        else:
            user_message = f"Error processing your question: {str(e)}"
            
        return {
            "content": user_message,
            "token_usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
        }

def batch_process_questions(questions):
    """Process multiple physics questions using standard API"""
    results = []
    for question in questions:
        if question.strip():  # Skip empty questions
            answer_data = process_physics_question(question)
            results.append({
                "question": question,
                "answer": answer_data["content"],
                "token_usage": answer_data["token_usage"]
            })
    return results

def create_pdf(results, filename):
    """Create a PDF document with questions and answers"""
    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    content = []
    
    # Title
    title_style = styles['Title']
    content.append(Paragraph("Physics Solutions", title_style))
    content.append(Spacer(1, 12))
    
    # Add each question and answer
    for i, result in enumerate(results, 1):
        # Question
        question_style = styles['Heading2']
        content.append(Paragraph(f"Question {i}:", question_style))
        content.append(Paragraph(result["question"], styles['BodyText']))
        content.append(Spacer(1, 12))
        
        # Answer
        content.append(Paragraph("Solution:", styles['Heading3']))
        
        # Split the answer by newlines to preserve formatting
        answer_lines = result["answer"].split('\n')
        for line in answer_lines:
            if line.strip():
                content.append(Paragraph(line, styles['BodyText']))
            else:
                content.append(Spacer(1, 6))
        
        # Add token usage if available
        if "token_usage" in result:
            content.append(Spacer(1, 12))
            content.append(Paragraph("Token Usage:", styles['Heading4']))
            token_info = f"Prompt Tokens: {result['token_usage']['prompt_tokens']} | " \
                         f"Completion Tokens: {result['token_usage']['completion_tokens']} | " \
                         f"Total Tokens: {result['token_usage']['total_tokens']}"
            content.append(Paragraph(token_info, styles['BodyText']))
        
        content.append(Spacer(1, 24))
    
    doc.build(content)
    return pdf_path

def create_docx(results, filename):
    """Create a DOCX document with questions and answers"""
    docx_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc = Document()
    
    # Add title
    doc.add_heading('Physics Solutions', 0)
    
    # Add each question and answer
    for i, result in enumerate(results, 1):
        doc.add_heading(f'Question {i}:', level=2)
        doc.add_paragraph(result["question"])
        
        doc.add_heading('Solution:', level=3)
        # Split the answer by newlines to preserve formatting
        answer_lines = result["answer"].split('\n')
        for line in answer_lines:
            if line.strip():
                doc.add_paragraph(line)
        
        # Add token usage if available
        if "token_usage" in result:
            doc.add_heading('Token Usage:', level=4)
            token_info = f"Prompt Tokens: {result['token_usage']['prompt_tokens']} | " \
                         f"Completion Tokens: {result['token_usage']['completion_tokens']} | " \
                         f"Total Tokens: {result['token_usage']['total_tokens']}"
            doc.add_paragraph(token_info)
    
    doc.save(docx_path)
    return docx_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api_test')
def api_test():
    """Test route to check OpenAI API connectivity"""
    try:
        # Just list models - a simple API call to check connectivity
        models = openai_client.models.list()
        return jsonify({
            "status": "success",
            "message": "Successfully connected to OpenAI API",
            "api_key_configured": bool(openai.api_key),
            "models_count": len(models.data) if hasattr(models, 'data') else 0,
            "proxy_enabled": USE_PROXY,
            "proxies": proxies if USE_PROXY else "Not using proxies"
        })
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Error connecting to OpenAI API: {str(e)}",
            "api_key_configured": bool(openai.api_key),
            "proxy_enabled": USE_PROXY,
            "proxies": proxies if USE_PROXY else "Not using proxies"
        }), 500

@app.route('/toggle_proxy')
def toggle_proxy():
    """Toggle proxy usage"""
    global USE_PROXY, proxies, openai_client
    
    # Toggle the proxy flag
    USE_PROXY = not USE_PROXY
    
    # Reconfigure environment
    if USE_PROXY:
        # Set proxies
        proxies = {
            "http://": http_proxy,
            "https://": https_proxy
        }
        # Set environment variables
        os.environ["HTTP_PROXY"] = http_proxy
        os.environ["HTTPS_PROXY"] = https_proxy
        print(f"Proxy enabled: {proxies}")
    else:
        # Clear proxies
        proxies = None
        # Unset environment variables
        if "HTTP_PROXY" in os.environ:
            del os.environ["HTTP_PROXY"]
        if "HTTPS_PROXY" in os.environ:
            del os.environ["HTTPS_PROXY"]
        print("Proxy disabled")
    
    # Recreate the client (timeout will be taken from environment)
    openai_client = OpenAI(
        api_key=openai.api_key,
        timeout=httpx.Timeout(60.0, connect=10.0)
    )
    
    return jsonify({
        "status": "success",
        "proxy_enabled": USE_PROXY,
        "message": f"Proxy {'enabled' if USE_PROXY else 'disabled'}"
    })

@app.route('/process_file', methods=['POST'])
def process_file():
    """Process a text file containing multiple questions"""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    if file and file.filename.endswith('.txt'):
        # Read questions from file
        questions = []
        for line in file:
            line_text = line.decode('utf-8').strip()
            if line_text:  # Skip empty lines
                questions.append(line_text)
        
        if not questions:
            return jsonify({"error": "No questions found in file"}), 400
        
        # Process questions in batch
        results = batch_process_questions(questions)
        return jsonify({"results": results})
    
    return jsonify({"error": "Invalid file format. Only .txt files are supported"}), 400

@app.route('/process', methods=['POST'])
def process_question():
    if request.method == 'POST':
        data = request.get_json()
        
        if 'questions' in data and isinstance(data['questions'], list):
            # Batch processing
            results = batch_process_questions(data['questions'])
            return jsonify({"results": results})
        
        elif 'question' in data:
            # Single question processing
            question = data['question']
            answer_data = process_physics_question(question)
            return jsonify({
                "question": question,
                "answer": answer_data["content"],
                "token_usage": answer_data["token_usage"]
            })
        
        return jsonify({"error": "Invalid request format"}), 400

@app.route('/process_image', methods=['POST'])
def process_image():
    if 'image' not in request.files:
        return jsonify({"error": "No image file provided"}), 400
    
    file = request.files['image']
    
    if file.filename == '':
        return jsonify({"error": "No image selected"}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Extract text from image
        extracted_text = extract_text_from_image(filepath)
        
        if not extracted_text:
            return jsonify({"error": "Could not extract text from image"}), 400
        
        # Process the extracted text
        answer_data = process_physics_question(extracted_text)
        
        # Clean up the file
        try:
            os.remove(filepath)
        except:
            pass
        
        return jsonify({
            "question": extracted_text,
            "answer": answer_data["content"],
            "token_usage": answer_data["token_usage"]
        })
    
    return jsonify({"error": "Invalid file format"}), 400

@app.route('/export', methods=['POST'])
def export_document():
    data = request.get_json()
    
    if 'results' not in data or not isinstance(data['results'], list):
        return jsonify({"error": "Invalid data format"}), 400
    
    file_format = data.get('format', 'pdf').lower()
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    
    if file_format == 'pdf':
        filename = f"physics_solutions_{timestamp}.pdf"
        file_path = create_pdf(data['results'], filename)
        return jsonify({"file_url": f"/download/{filename}"})
    
    elif file_format == 'docx':
        filename = f"physics_solutions_{timestamp}.docx"
        file_path = create_docx(data['results'], filename)
        return jsonify({"file_url": f"/download/{filename}"})
    
    else:
        return jsonify({"error": "Unsupported format"}), 400

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    return send_file(
        os.path.join(app.config['UPLOAD_FOLDER'], filename),
        as_attachment=True,
        download_name=filename
    )

if __name__ == '__main__':
    app.run(debug=True)
